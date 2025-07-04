from fastapi import APIRouter, Depends, HTTPException, status, UploadFile, File, Form, Request
from fastapi.responses import StreamingResponse, FileResponse, HTMLResponse, JSONResponse
from sqlalchemy.orm import Session
from typing import List, Dict, Optional
from ..database import get_db
from ..models.user import User
from ..models.chat import ChatMessage, FileOutput
from ..models.agent import Agent
from ..models.model_settings import ModelSettings
from ..models.api_key import APIKey
from ..models.vector_source import VectorSource
from ..schemas.chat import (
    ChatMessageCreate, 
    ChatMessageResponse, 
    ChatHistoryResponse,
    ConnectedSource,
    FileAttachmentResponse
)
from ..utils.auth import get_current_user, create_access_token
from ..utils.ai_client import get_ai_response_from_model, get_ai_response_from_vectorstore
from ..services.vector_service import VectorService
from ..utils.api_key_validator import validate_finiite_api_key
import os
import uuid
import shutil
import json
import base64
import csv
import io
from ..models.chat import FileAttachment
from fpdf import FPDF
from docx import Document
from ..utils.activity_logger import log_activity
from ..services.trial_service import TrialService
from datetime import datetime, timedelta
from ..config import config

router = APIRouter(prefix="/chat", tags=["Chat"])

# Define upload directory and create it if it doesn't exist
UPLOAD_DIR = "uploads"
if not os.path.exists(UPLOAD_DIR):
    os.makedirs(UPLOAD_DIR)

async def save_uploaded_file(file: UploadFile, user_id: int) -> Dict:
    """Save uploaded file and return metadata"""
    file_ext = file.filename.split(".")[-1].lower()
    file_id = str(uuid.uuid4())
    file_path = os.path.join(UPLOAD_DIR, f"{user_id}_{file_id}.{file_ext}")
    
    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
    
    return {
        "name": file.filename,
        "type": file_ext,
        "url": file_path,
        "size": os.path.getsize(file_path)
    }
        

@router.post("/{agent_id}/messages", response_model=ChatMessageResponse)
async def create_message(
    agent_id: int,
    content: str = Form(...),
    model: str = Form(...),
    files: List[UploadFile] = File(default=[]),
    source_ids: str = Form(default="[]"),
    current_user: User = Depends(get_current_user),
    request: Request = None,
    db: Session = Depends(get_db)
):
    """Create a new chat message and get AI response"""
    try:
        selected_source_ids = json.loads(source_ids)
    except json.JSONDecodeError:
        selected_source_ids = []

    # Check if agent exists and belongs to user
    agent = db.query(Agent).filter(
        Agent.id == agent_id,
        Agent.user_id == current_user.id
    ).first()
    if not agent:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Agent not found"
        )

    # Get available vector sources
    if selected_source_ids:
        available_sources = db.query(VectorSource).filter(
            VectorSource.user_id == current_user.id,
            VectorSource.id.in_(selected_source_ids)
        ).all()
    else:
        available_sources = db.query(VectorSource).filter(
            VectorSource.user_id == current_user.id,
            VectorSource.id.in_(agent.vector_sources_ids or [])
        ).all()

    # Check if requested model is enabled
    model_setting = db.query(ModelSettings).filter(
        ModelSettings.user_id == current_user.id,
        ModelSettings.ai_model_name == model,
        ModelSettings.is_enabled == True
    ).first()
    if not model_setting:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Model {model} is not enabled or not found"
        )
    # Check API key for the model's provider
    api_key = db.query(APIKey).filter(
        APIKey.user_id == current_user.id,
        APIKey.provider == model_setting.provider,
        APIKey.is_valid == True
    ).first()
    openai_api_key = db.query(APIKey).filter(
        APIKey.user_id == current_user.id,
        APIKey.provider == "openai",
        APIKey.is_valid == True
    ).first()
    if not api_key:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"No valid API key found for OpenAI"
        )

    # Check daily message limit for trial users
    today_start = datetime.now().replace(hour=0, minute=0, second=0, microsecond=0)
    today_messages = db.query(ChatMessage).filter(
        ChatMessage.user_id == current_user.id,
        ChatMessage.created_at >= today_start
    ).count()
    
    TrialService.check_trial_limits(db, current_user, 'messages_per_day', today_messages)
    print(model)
    # Save user message
    user_message = ChatMessage(
        agent_id=agent_id,
        user_id=current_user.id,
        role="user",
        content=content,
        model=model
    )
    db.add(user_message)
    db.commit()
    db.refresh(user_message)

    # Log chat activity
    await log_activity(
        db=db,
        user_id=current_user.id,
        activity_type="chat_message",
        description=f"Sent message to agent: {agent.name}",
        request=request,
        metadata={
            "agent_id": agent_id,
            "agent_name": agent.name,
            "message_id": user_message.id,
            "model": model,
            "has_files": len(files) > 0,
            "sources_count": len(available_sources)
        }
    )

    # Save file attachments
    file_attachments = []
    for file in files:
        try:
            attachment_data = await save_uploaded_file(file, current_user.id)
            # Create FileAttachment record
            file_attachment = FileAttachment(
                message_id=user_message.id,
                name=attachment_data["name"],
                type=attachment_data["type"],
                url=attachment_data["url"],
                size=attachment_data["size"]
            )
            db.add(file_attachment)
            file_attachments.append(attachment_data)
        except Exception as e:
            db.rollback()
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Error saving file: {str(e)}"
            )
    
    db.commit()

    try:
        # Get chat history
        chat_history = db.query(ChatMessage).filter(
            ChatMessage.agent_id == agent_id,
            ChatMessage.user_id == current_user.id
        ).order_by(ChatMessage.created_at.asc()).all()
        
        # Format chat history into messages
        formatted_messages = []
        for msg in chat_history:
            formatted_messages.append({
                "role": msg.role,
                "content": msg.content
            })
        
        # Add current message
            formatted_messages.append({
                "role": "user",
                "content": content
            })
        
        # Initialize VectorService and search similar content
        vector_service = VectorService(current_user.id)
        
        # Prepare the final response content and references
        response_content = ""
        similar_results = []

        # Search through connected vector sources only
        for vector_source in available_sources:
            try:
                results = await vector_service.search_similar(
                    query=content,
                    source_name=vector_source.table_name,
                    embedding_model=vector_source.embedding_model,
                    api_key=openai_api_key.api_key
                )
                # Add source information to results
                for result in results:
                    result['source_name'] = vector_source.name
                    result['table_name'] = vector_source.table_name
                    result['is_connected'] = True  # All sources are now connected by definition
                similar_results.extend(results)
            except Exception as e:
                print(f"Error searching vector source {vector_source.name}: {str(e)}")
                continue

        # Format the response with similar content if results found
        if similar_results:
            # Sort results by relevance score
            similar_results.sort(key=lambda x: x.get('score', 0), reverse=True)
            
            message_from_vector = ""
            connected_sources = []
            # Create a mapping of table_name to source info
            source_mapping = {
                source.table_name: {
                    "id": source.id,
                    "name": source.name,
                    "type": source.source_type,
                    "connection_settings": source.connection_settings
                } for source in available_sources
            }
            
            for result in similar_results:
                message_from_vector += f"[From {result['source_name']}]: {result['content']}\n"
                source_info = source_mapping.get(result['table_name'])
                if source_info and source_info["id"] not in [s.get("id") for s in connected_sources]:
                    connected_sources.append(source_info)
            
            conversation = {
                "messages": message_from_vector,
                "agent_instructions": agent.instructions,
                "model": model,
                "provider": model_setting.provider,
                "api_key": api_key.api_key,
                "attachments": file_attachments,
                "query": content,
                "reference_enabled": agent.reference_enabled
            }
            response_content = await get_ai_response_from_vectorstore(conversation)

            # Add connected sources to the response
            response_content = {
                "content": response_content,
                "connected_sources": connected_sources if agent.reference_enabled else []
            }
        
        # If no results found in vector search
        if not response_content:
            conversation = {
                "messages": formatted_messages,
                "agent_instructions": agent.instructions,
                "model": model,
                "provider": model_setting.provider,
                "api_key": api_key.api_key,
                "attachments": file_attachments
            }
            response_content = await get_ai_response_from_model(conversation)
            response_content = {
                "content": response_content,
                "connected_sources": []
            }

        # Check if this is a file generation request
        file_keywords = ["provide", "generate", "create", "download", "export", "save", "convert", "give me", "create a", "generate a"]
        file_types = ["csv", "pdf", "doc", "docx", "document"]
        
        is_file_request = (
            any(keyword in content.lower() for keyword in file_keywords) and
            any(ftype in content.lower() for ftype in file_types)
        )
        
        if is_file_request:
            # Determine file type based on content
            file_type = None
            if "csv" in content.lower():
                file_type = "csv"
            elif "pdf" in content.lower():
                file_type = "pdf"
            elif any(doc_type in content.lower() for doc_type in ["doc", "docx", "document"]):
                file_type = "doc"
            
            # Extract the actual request from the user's message
            # Remove file type and generation keywords to get the core request
            core_request = content.lower()
            for keyword in file_keywords + file_types:
                core_request = core_request.replace(keyword, "")
            core_request = core_request.strip()
            
            # Get AI response for file content based on file type
            if file_type == "pdf":
                agent_instructions = (
                    f"You are tasked with generating a PDF document about: {core_request}\n\n"
                    "Follow these guidelines:\n"
                    "1. Focus ONLY on the specific information requested\n"
                    "2. Format the content in a clear, structured way suitable for PDF\n"
                    "3. Include a title that reflects the specific request\n"
                    "4. Organize content with appropriate headers and sections\n"
                    "5. Return ONLY the content, no explanations or markdown\n"
                    "6. Ensure proper spacing between sections\n"
                    "7. Keep the formatting simple and compatible with PDF generation\n"
                    "8. Start with the most relevant information first\n"
                    "9. Use clear section headers ending with ':' for better formatting"
                )
            elif file_type == "doc":
                agent_instructions = (
                    f"You are tasked with generating a Word document about: {core_request}\n\n"
                    "Follow these guidelines:\n"
                    "1. Focus ONLY on the specific information requested\n"
                    "2. Format the content in a clear, structured way suitable for a document\n"
                    "3. Include a title that reflects the specific request\n"
                    "4. Organize content with appropriate headers and sections\n"
                    "5. Return ONLY the content, no explanations or markdown\n"
                    "6. Ensure proper spacing between sections\n"
                    "7. Keep the formatting simple and compatible with document generation\n"
                    "8. Start with the most relevant information first\n"
                    "9. Use clear section headers ending with ':' for better formatting"
                )
            else:  # CSV
                agent_instructions = (
                    f"You are tasked with generating CSV data about: {core_request}\n\n"
                    "Follow these guidelines:\n"
                    "1. Focus ONLY on the specific information requested\n"
                    "2. Return the data in this format:\n"
                    "   ```json\n"
                    "   [\n"
                    "     {\n"
                    "       \"column1\": \"value1\",\n"
                    "       \"column2\": \"value2\"\n"
                    "     }\n"
                    "   ]\n"
                    "   ```\n"
                    "3. Use clear, descriptive column names\n"
                    "4. Include all relevant data fields\n"
                    "5. Ensure data is properly structured and consistent\n"
                    "6. Return ONLY the JSON array, no explanations\n"
                    "7. Make sure all records follow the same schema\n"
                    "8. Use proper data types for each field"
                )
            
            # Add the original user message for context
            formatted_messages.append({
                "role": "user",
                "content": f"Original request: {content}\nCore request: {core_request}"
            })
            
            conversation = {
                "messages": formatted_messages,
                "agent_instructions": agent_instructions,
                "model": model,
                "provider": model_setting.provider,
                "api_key": api_key.api_key
            }
            
            file_content = await get_ai_response_from_model(conversation)
            
            # Convert to PDF or DOC if needed
            clean_content = extract_data_only(file_content, file_type)
            if file_type == "pdf":
                pdf = FPDF()
                pdf.set_auto_page_break(auto=True, margin=15)
                pdf.add_page()
                
                # Extract title from content or use a default
                title = core_request.title()
                content_lines = clean_content.split('\n')
                if content_lines and not content_lines[0].endswith(':'):
                    title = content_lines[0]
                    content_lines = content_lines[1:]
                
                # Add title with larger font
                pdf.set_font("Arial", 'B', 14)
                pdf.cell(0, 10, title, ln=True, align='C')
                pdf.ln(10)
                
                # Add content with regular font
                pdf.set_font("Arial", size=12)
                for line in content_lines:
                    if line.strip():  # Skip empty lines
                        if line.strip().endswith(':'):  # Treat as section header
                            pdf.ln(5)
                            pdf.set_font("Arial", 'B', 12)
                            pdf.cell(0, 10, line, ln=True)
                            pdf.set_font("Arial", size=12)
                        else:
                            pdf.multi_cell(0, 8, line)
                
                pdf_bytes = pdf.output(dest='S').encode('latin1')
                file_content_to_save = base64.b64encode(pdf_bytes).decode('utf-8')
            
            elif file_type == "doc":
                doc = Document()
                
                # Extract title from content or use a default
                title = core_request.title()
                content_lines = clean_content.split('\n')
                if content_lines and not content_lines[0].endswith(':'):
                    title = content_lines[0]
                    content_lines = content_lines[1:]
                
                # Add title
                doc.add_heading(title, level=1)
                
                # Add content with proper formatting
                for line in content_lines:
                    if line.strip():  # Skip empty lines
                        if line.strip().endswith(':'):  # Treat as section header
                            doc.add_heading(line, level=2)
                        else:
                            doc.add_paragraph(line)
                
                doc_io = io.BytesIO()
                doc.save(doc_io)
                doc_bytes = doc_io.getvalue()
                file_content_to_save = base64.b64encode(doc_bytes).decode('utf-8')
            
            else:  # CSV
                try:
                    # Try to parse as JSON first (in case AI returned structured data)
                    try:
                        data = json.loads(clean_content)
                        if isinstance(data, list) and len(data) > 0:
                            # If it's a list of dictionaries, convert to CSV
                            output = io.StringIO()
                            writer = csv.DictWriter(output, fieldnames=data[0].keys())
                            writer.writeheader()
                            writer.writerows(data)
                            file_content_to_save = output.getvalue()
                        else:
                            file_content_to_save = clean_content
                    except json.JSONDecodeError:
                        # If not JSON, use the cleaned content directly
                        file_content_to_save = clean_content
                except Exception as e:
                    raise HTTPException(
                        status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                        detail=f"Error processing CSV content: {str(e)}"
                    )
            
            # Create a unique filename
            file_name = f"generated_{uuid.uuid4().hex[:8]}.{file_type}"
            
            # Create FileOutput record
            file_output = FileOutput(
                message_id=user_message.id,
                name=file_name,
                type=file_type,
                content=file_content_to_save
            )
            db.add(file_output)
            db.commit()
            db.refresh(file_output)
            
            # Create assistant message with download link
            response_content = f"I've generated the {file_type.upper()} file for you. You can download it using this link: [Download {file_name}](/download/{file_output.id})"

        ai_message = ChatMessage(
            agent_id=agent_id,
            user_id=current_user.id,
            role="assistant",
            content=json.dumps({
                "content": response_content,
                "connected_sources": []
            }) if isinstance(response_content, str) else json.dumps(response_content),  # Store the complete content as JSON
            model=model
        )
        db.add(ai_message)
        db.commit()
        db.refresh(ai_message)

        # Format the response using ChatMessageResponse model
        try:
            content_data = json.loads(ai_message.content)
            response_content = content_data.get("content", ai_message.content)
            connected_sources = content_data.get("connected_sources", [])
        except (json.JSONDecodeError, AttributeError):
            response_content = ai_message.content
            connected_sources = []

        return ChatMessageResponse(
            id=ai_message.id,
            agent_id=ai_message.agent_id,
            user_id=ai_message.user_id,
            role=ai_message.role,
            content=response_content,
            model=ai_message.model,
            created_at=ai_message.created_at,
            updated_at=ai_message.updated_at,
            attachments=[],  # No attachments for AI response
            connected_sources=connected_sources,
            citations=[],
            search_results=[],
            choices=[]
        )

    except Exception as e:
        db.rollback()
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error getting AI response: {str(e)}"
        )

@router.get("/{agent_id}/history", response_model=ChatHistoryResponse)
async def get_chat_history(
    agent_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get chat history for an agent"""
    # Query messages with attachments using joinedload
    messages = db.query(ChatMessage).filter(
        ChatMessage.agent_id == agent_id,
        ChatMessage.user_id == current_user.id
    ).order_by(ChatMessage.created_at.asc()).all()
    
    # Get all vector sources for mapping
    vector_sources = db.query(VectorSource).filter(
        VectorSource.user_id == current_user.id
    ).all()
    source_mapping = {
        source.name: {
            "id": source.id,
            "name": source.name,
            "type": source.source_type
        } for source in vector_sources
    }
    
    # Format response with attachments
    formatted_messages = []
    for msg in messages:
        # Query attachments for this message
        attachments = db.query(FileAttachment).filter(
            FileAttachment.message_id == msg.id
        ).all()
        
        # Convert attachments to dictionaries
        attachment_dicts = [{
            "id": att.id,
            "name": att.name,
            "type": att.type,
            "url": att.url,
            "size": att.size
        } for att in attachments]
        
        # Initialize metadata
        content = msg.content
        connected_sources = []
        search_metadata = {
            "citations": [],
            "search_results": [],
            "choices": []
        }
        
        # Parse content if it's JSON
        if msg.role == "assistant":
            try:
                content_data = json.loads(msg.content)
                if isinstance(content_data, dict):
                    content = content_data.get("content", msg.content)
                    # Handle connected sources
                    sources_data = content_data.get("connected_sources", [])
                    if sources_data:
                        if isinstance(sources_data[0], str) and sources_data[0].endswith('.pdf'):
                            # Convert old format (source names) to new format (source info)
                            connected_sources = [
                                source_mapping[name]
                                for name in sources_data
                                if name in source_mapping
                            ]
                        elif isinstance(sources_data[0], dict):
                            # Already in new format (source info)
                            connected_sources = sources_data
                        else:
                            # Handle case where it's a list of IDs
                            source_id_to_info = {
                                s.id: {
                                    "id": s.id,
                                    "name": s.name,
                                    "type": s.source_type,
                                    "connection_settings": s.connection_settings
                                } for s in vector_sources
                            }
                            connected_sources = [
                                source_id_to_info[sid]
                                for sid in sources_data
                                if sid in source_id_to_info
                            ]
                    
                    # Handle web search metadata if present
                    if "search_metadata" in content_data:
                        metadata = content_data["search_metadata"]
                        search_metadata.update({
                            "citations": metadata.get("citations", []),
                            "search_results": metadata.get("search_results", []),
                            "choices": metadata.get("choices", [])
                        })
            except (json.JSONDecodeError, IndexError, KeyError, TypeError):
                pass
        
        formatted_message = ChatMessageResponse(
            id=msg.id,
            agent_id=msg.agent_id,
            user_id=msg.user_id,
            role=msg.role,
            content=content,
            model=msg.model,
            created_at=msg.created_at,
            updated_at=msg.updated_at,
            attachments=attachment_dicts,
            connected_sources=[ConnectedSource(**source) for source in connected_sources],
            citations=search_metadata["citations"],
            search_results=search_metadata["search_results"],
            choices=search_metadata["choices"]
        )
        
        formatted_messages.append(formatted_message)
    
    return ChatHistoryResponse(messages=formatted_messages)

@router.delete("/{agent_id}/history", status_code=status.HTTP_204_NO_CONTENT)
async def clear_chat_history(
    agent_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Clear chat history for an agent"""
    db.query(ChatMessage).filter(
        ChatMessage.agent_id == agent_id,
        ChatMessage.user_id == current_user.id
    ).delete()
    db.commit()

@router.post("/{agent_id}/web-search", response_model=ChatMessageResponse)
async def web_search(
    agent_id: int,
    content: str = Form(...),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Perform web search using Perplexity AI"""
    try:
        import httpx
        import os
        
        # Save user query message first
        user_message = ChatMessage(
            agent_id=agent_id,
            user_id=current_user.id,
            role="user",
            content=f"[Web Search Query] {content}",
            model="sonar"  # Updated model name
        )
        db.add(user_message)
        db.commit()

        # Get Perplexity API key from environment
        api_key = os.getenv("PERPLEXITY_API_KEY")
        if not api_key:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail="Perplexity API key not configured"
            )

        # Make request to Perplexity API
        try:
            async with httpx.AsyncClient() as client:
                request_data = {
                    "model": "sonar",
                    "messages": [
                        {"role": "system", "content": "Be precise and concise."},
                        {"role": "user", "content": content}
                    ]
                }

                response = await client.post(
                    "https://api.perplexity.ai/chat/completions",
                    headers={
                        "Content-Type": "application/json",
                        "Authorization": f"Bearer {api_key}"
                    },
                    json=request_data,
                    timeout=30.0  # Add explicit timeout
                )
                
                if response.status_code != 200:
                    error_detail = f"Perplexity API error (Status {response.status_code}): {response.text}"
                    raise HTTPException(
                        status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                        detail=error_detail
                    )

                result = response.json()

        except httpx.RequestError as e:
            error_msg = f"Failed to connect to Perplexity API: {str(e)}"
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=error_msg
            )
        except json.JSONDecodeError as e:
            error_msg = f"Invalid JSON response from Perplexity API: {str(e)}"
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=error_msg
            )
        except Exception as e:
            error_msg = f"Unexpected error during API request: {str(e)}"
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=error_msg
            )
            
        # Format the response content with citations and search results
        try:
            ai_content = result["choices"][0]["message"]["content"]
        except (KeyError, IndexError) as e:
            error_msg = f"Unexpected response format from Perplexity API: {str(e)}"
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=error_msg
            )
            
        # Store additional search information in the message
        search_metadata = {
            "citations": result.get("citations", []),
            "search_results": result.get("search_results", []),
            "choices": result.get("choices", [])
        }
            
        # Create a formatted message that includes both the AI response and metadata
        full_content = {
            "content": ai_content,
            "connected_sources": [],  # Web search doesn't use connected sources
            "search_metadata": search_metadata
        }

        # Save the enhanced response as assistant's message with complete search data
        ai_message = ChatMessage(
            agent_id=agent_id,
            user_id=current_user.id,
            role="assistant",
            content=json.dumps(full_content),  # Store the complete content as JSON
            model="sonar"
        )
        db.add(ai_message)
        db.commit()
        db.refresh(ai_message)

        # Format the response using ChatMessageResponse model
        content_data = json.loads(ai_message.content)
        return ChatMessageResponse(
            id=ai_message.id,
            agent_id=ai_message.agent_id,
            user_id=ai_message.user_id,
            role=ai_message.role,
            content=content_data.get("content", ai_message.content),
            model=ai_message.model,
            created_at=ai_message.created_at,
            updated_at=ai_message.updated_at,
            attachments=[],  # No attachments for web search response
            connected_sources=content_data.get("connected_sources", []),
            citations=search_metadata.get("citations", []),
            search_results=search_metadata.get("search_results", []),
            choices=search_metadata.get("choices", [])
        )

    except Exception as e:
        db.rollback()
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error performing web search: {str(e)}"
        ) 

@router.get("/download/{file_id}")
async def download_file(
    file_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Download a generated file"""
    # Get the file output record
    file_output = db.query(FileOutput).filter(FileOutput.id == file_id).first()
    if not file_output:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="File not found"
        )
    
    # Verify the user has access to this file
    message = db.query(ChatMessage).filter(ChatMessage.id == file_output.message_id).first()
    if not message or message.user_id != current_user.id:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Access denied"
        )
    
    try:
        if file_output.type == "csv":
            # Parse the CSV content and create a file-like object
            content = io.StringIO()
            if file_output.content.startswith("[") and file_output.content.endswith("]"):
                # Handle JSON array format
                rows = json.loads(file_output.content)
                if rows:
                    writer = csv.DictWriter(content, fieldnames=rows[0].keys())
                    writer.writeheader()
                    writer.writerows(rows)
            else:
                # Handle raw CSV content
                content.write(file_output.content)
            
            content.seek(0)
            return StreamingResponse(
                iter([content.getvalue()]),
                media_type="text/csv",
                headers={"Content-Disposition": f"attachment; filename={file_output.name}"}
            )
            
        elif file_output.type in ["pdf", "doc"]:
            try:
                content = base64.b64decode(file_output.content)
                content_io = io.BytesIO(content)
                media_type = "application/pdf" if file_output.type == "pdf" else "application/msword"
                return StreamingResponse(
                    iter([content_io.getvalue()]),
                    media_type=media_type,
                    headers={"Content-Disposition": f"attachment; filename={file_output.name}"}
                )
            except Exception as e:
                raise HTTPException(
                    status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    detail=f"Error processing file content: {str(e)}"
                )
            
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error generating file: {str(e)}"
        ) 

def extract_data_only(text, file_type="csv"):
    """
    Extract clean content from AI response based on file type.
    Args:
        text (str): The raw text from AI response
        file_type (str): The type of file being generated (csv, pdf, doc)
    Returns:
        str: Cleaned content suitable for the specified file type
    """
    import re
    import json
    
    # First try to extract content from code blocks
    code_block_pattern = r'```(?:json|python)?\n([\s\S]*?)```'
    code_blocks = re.findall(code_block_pattern, text)
    
    if code_blocks:
        # If we found code blocks, use the first one that contains valid data
        for block in code_blocks:
            if file_type == "csv":
                # For CSV, try to parse as JSON first
                try:
                    data = json.loads(block)
                    if isinstance(data, list) and len(data) > 0:
                        return json.dumps(data)  # Return JSON string for later processing
                except json.JSONDecodeError:
                    # If not JSON, check if it looks like CSV
                    if ',' in block and '\n' in block:
                        return block.strip()
            else:
                # For PDF/DOC, look for text content
                if block.strip() and not block.startswith('import '):
                    return block.strip()
    
    # If no suitable code blocks found, process the entire text
    if file_type == "csv":
        # Try to find CSV-like content (lines with commas)
        lines = text.splitlines()
        csv_lines = []
        for line in lines:
            # Skip empty lines and lines that look like explanations
            if line.strip() and ',' in line and not line.startswith(('#', '//', '--', '```')):
                csv_lines.append(line.strip())
        return '\n'.join(csv_lines)
    
    else:  # pdf or doc
        # Remove markdown and code formatting
        text = re.sub(r'```[a-zA-Z]*\n[\s\S]*?```', '', text)  # Remove code blocks
        text = re.sub(r'[`*_]', '', text)  # Remove markdown characters
        text = re.sub(r'^#+\s+', '', text, flags=re.MULTILINE)  # Remove markdown headers
        text = re.sub(r'\n{3,}', '\n\n', text)  # Normalize multiple newlines
        text = re.sub(r'^\s*[-+*]\s', '', text, flags=re.MULTILINE)  # Remove list markers
        
        # Clean up the text
        lines = text.splitlines()
        clean_lines = []
        for line in lines:
            line = line.strip()
            if line and not line.startswith(('//', '--', '#', '```')):
                clean_lines.append(line)
        
        return '\n'.join(clean_lines) 

@router.get("/widget/{agent_id}/config")
async def get_widget_config(
    agent_id: int,
    request: Request,
    db: Session = Depends(get_db)
):
    """Get widget configuration for the specified agent"""
    # Get agent details
    agent = db.query(Agent).filter(Agent.id == agent_id).first()
    if not agent:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Agent not found"
        )

    # Get the origin domain
    origin = request.headers.get("origin", "")
    
    # Check if the origin is allowed (you may want to implement more strict checks)
    if origin not in origins:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Origin not allowed"
        )

    return {
        "agent_id": agent.id,
        "name": agent.name,
        "greeting": agent.greeting or "Hello! How can I help you today?",
        "theme": agent.theme or "light",
        "position": "bottom-right",  # Default position
        "height": "600px",
        "width": "400px"
    }

@router.get("/widget/{agent_id}/embed")
async def get_widget_embed(
    agent_id: int,
    request: Request,
    db: Session = Depends(get_db)
):
    """Get the HTML embed code for the chat widget"""
    # Get agent details
    agent = db.query(Agent).filter(Agent.id == agent_id).first()
    if not agent:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Agent not found"
        )

    # Base URL for the widget
    base_url = str(request.base_url).rstrip('/')
    
    # Generate the widget HTML
    widget_html = f"""
    <div id="finiite-chat-widget"></div>
    <script>
        (function() {{
            var script = document.createElement('script');
            script.src = '{base_url}/static/widget.js';
            script.async = true;
            script.onload = function() {{
                initFiniiteWidget({{
                    agentId: '{agent_id}',
                    baseUrl: '{base_url}'
                }});
            }};
            document.head.appendChild(script);
        }})();
    </script>
    """
    
    return HTMLResponse(content=widget_html) 

@router.post("/embed/{agent_id}/messages", response_model=ChatMessageResponse)
async def create_message_with_api_key(
    agent_id: int,
    content: str = Form(...),
    model: str = Form(...),
    api_key: str = Form(...),
    files: List[UploadFile] = File(default=[]),
    source_ids: str = Form(default="[]"),
    request: Request = None,
    db: Session = Depends(get_db)
):
    """Create a new chat message using Finiite API key authentication"""
    # Validate Finiite API key
    if not await validate_finiite_api_key(api_key):
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Invalid Finiite API key"
        )
    
    # Get user by API key
    user = db.query(User).filter(
        User.finiite_api_key == api_key
    ).first()
    if not user:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="User not found"
        )

    # Generate access token for the user
    access_token_expires = timedelta(minutes=int(config["ACCESS_TOKEN_EXPIRE_MINUTES"]))
    access_token = create_access_token(
        data={"sub": user.email}, expires_delta=access_token_expires
    )

    # Create message using the existing endpoint logic
    return await create_message(
        agent_id=agent_id,
        content=content,
        model=model,
        files=files,
        source_ids=source_ids,
        current_user=user,
        request=request,
        db=db
    )

@router.post("/embed/{agent_id}/history", response_model=ChatHistoryResponse)
async def get_chat_history_with_api_key(
    agent_id: int,
    api_key: str = Form(...),
    db: Session = Depends(get_db)
):
    """Get chat history using Finiite API key authentication"""
    # Validate Finiite API key
    if not await validate_finiite_api_key(api_key):
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Invalid Finiite API key"
        )
    
    # Get user by API key
    user = db.query(User).filter(
        User.finiite_api_key == api_key
    ).first()
    if not user:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="User not found"
        )

    # Get chat history using the existing endpoint logic
    return await get_chat_history(
        agent_id=agent_id,
        current_user=user,
        db=db
    )

@router.post("/embed/download/{file_id}")
async def download_file_with_api_key(
    file_id: int,
    api_key: str = Form(...),
    db: Session = Depends(get_db)
):
    """Download a file using Finiite API key authentication"""
    # Validate Finiite API key
    if not await validate_finiite_api_key(api_key):
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Invalid Finiite API key"
        )
    
    # Get user by API key
    user = db.query(User).filter(
        User.finiite_api_key == api_key
    ).first()
    if not user:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="User not found"
        )

    # Download file using the existing endpoint logic
    return await download_file(
        file_id=file_id,
        current_user=user,
        db=db
    )

@router.post("/embed/{agent_id}/web-search", response_model=ChatMessageResponse)
async def web_search_with_api_key(
    agent_id: int,
    content: str = Form(...),
    api_key: str = Form(...),
    db: Session = Depends(get_db)
):
    """Perform web search using Finiite API key authentication"""
    # Validate Finiite API key
    if not await validate_finiite_api_key(api_key):
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Invalid Finiite API key"
        )
    
    # Get user by API key
    user = db.query(User).filter(
        User.finiite_api_key == api_key
    ).first()
    if not user:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="User not found"
        )

    # Perform web search using the existing endpoint logic
    return await web_search(
        agent_id=agent_id,
        content=content,
        current_user=user,
        db=db
    ) 
